from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Dosya yolları (mutlak yol kullanıyoruz)
BASE_DIR = r"D:\Projects\SelcukAiAssistant\docs\vize_raporu"
SOURCE_FILE = os.path.join(BASE_DIR, "VIZE_RAPORU. docx")
TEMPLATE_FILE = os.path.join(BASE_DIR, "sablon", "proje_sablonu. docx")
OUTPUT_FILE = os.path.join(BASE_DIR, "VIZE_RAPORU_FINAL.docx")

def check_files():
    """Dosyaların varlığını kontrol et"""
    print("🔍 Dosya kontrolü:")
    print(f"  Kaynak:  {SOURCE_FILE}")
    print(f"  Exists: {os.path.exists(SOURCE_FILE)}")
    print(f"  Şablon: {TEMPLATE_FILE}")
    print(f"  Exists: {os.path.exists(TEMPLATE_FILE)}")
    
    if not os.path.exists(SOURCE_FILE):
        print(f"\n❌ HATA: {SOURCE_FILE} bulunamadı!")
        print("\n📋 Mevcut dosyalar:")
        for f in os.listdir(BASE_DIR):
            print(f"   - {f}")
        return False
    
    if not os. path.exists(TEMPLATE_FILE):
        print(f"\n❌ HATA: {TEMPLATE_FILE} bulunamadı!")
        # Şablon yoksa mevcut dosyayı kopyala
        print("\n⚠️  Şablon bulunamadı, kaynak dosya üzerinden çalışılacak...")
        return "use_source"
    
    return True

def extract_sections(source_doc):
    """Kaynak dosyadan bölümleri çıkar"""
    doc = Document(source_doc)
    sections = {}
    current_section = None
    current_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Bölüm başlıklarını tespit et
        if text in ['ÖZET', 'ABSTRACT', 'ÖNSÖZ', 'GİRİŞ', 'KAYNAK ARAŞTIRMASI', 
                    'MATERYAL VE YÖNTEM', 'ARAŞTIRMA BULGULARI VE TARTIŞMA', 
                    'SONUÇLAR VE ÖNERİLER', 'KAYNAKLAR']:
            if current_section and current_content:
                sections[current_section] = '\n'. join(current_content)
            current_section = text
            current_content = []
        elif text.startswith('1. GİRİŞ') or text.startswith('2. KAYNAK'):
            if current_section and current_content:
                sections[current_section] = '\n'.join(current_content)
            current_section = text. split('.', 1)[1].strip() if '.' in text else text
            current_content = []
        elif current_section and text:
            current_content.append(text)
    
    # Son bölümü ekle
    if current_section and current_content:
        sections[current_section] = '\n'. join(current_content)
    
    return sections

def create_formatted_doc(sections, output_file):
    """Yeni formatlanmış doküman oluştur"""
    doc = Document()
    
    # Sayfa düzeni
    section = doc.sections[0]
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    
    # Stil tanımlamaları
    style = doc.styles['Normal']
    font = style.font
    font. name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    
    # İçeriği ekle
    for section_name, content in sections.items():
        # Bölüm başlığı
        heading = doc.add_paragraph(section_name)
        heading.style = 'Heading 1'
        heading.runs[0].font.name = 'Times New Roman'
        heading. runs[0].font.size = Pt(12)
        heading.runs[0].bold = True
        
        # Bölüm içeriği
        doc.add_paragraph(content)
        doc.add_paragraph()  # Boşluk
    
    doc.save(output_file)
    print(f"✅ Yeni doküman oluşturuldu: {output_file}")

def simple_fill_template(template_file, sections, output_file):
    """Şablonu doldur (basit versiyon)"""
    doc = Document(template_file)
    
    # Placeholder değiştirmeleri
    replacements = {
        'PROJE BAŞLIĞINI BURAYA YAZINIZ': 'SELÇUK ÜNİVERSİTESİ İÇİN YAPAY ZEKA DESTEKLİ AKADEMİK ASİSTAN',
        'Öğrencinin Adı SOYADI': 'Doğukan Balaman & Ali Yıldırım',
        'Unvanı Adı SOYADI': '[Danışman Unvanı Adı]',
        'Ay-Yıl':  'Aralık 2025',
        'Yıl, … Sayfa': '2025, 85 Sayfa',
        'Özet metnini yazmaya buradan başlayınız':  sections. get('ÖZET', ''),
        'Türkçe özet metninin İngilizce\'sini yazmaya buradan başlayınız': sections.get('ABSTRACT', ''),
        'Giriş bölümünü yazmaya buradan başlayınız': sections.get('GİRİŞ', ''),
    }
    
    # Tüm paragrafları tara
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                # Metni değiştir ama formatı koru
                for run in para.runs:
                    run.text = run.text. replace(old, new)
    
    doc.save(output_file)
    print(f"✅ Şablon dolduruldu: {output_file}")

if __name__ == "__main__": 
    print("=" * 60)
    print("📄 SELÇUK ÜNİVERSİTESİ RAPOR FORMATLAMA ARACI")
    print("=" * 60)
    
    # Dosya kontrolü
    file_check = check_files()
    
    if file_check == False:
        print("\n❌ Dosyalar bulunamadı.  Lütfen dosya yollarını kontrol edin.")
        exit(1)
    
    try:
        print("\n📖 İçerik çıkarılıyor...")
        sections = extract_sections(SOURCE_FILE)
        print(f"✅ {len(sections)} bölüm bulundu")
        
        if file_check == "use_source": 
            print("\n📝 Yeni formatlanmış doküman oluşturuluyor...")
            create_formatted_doc(sections, OUTPUT_FILE)
        else:
            print("\n📝 Şablon dolduruluyor...")
            simple_fill_template(TEMPLATE_FILE, sections, OUTPUT_FILE)
        
        print("\n" + "=" * 60)
        print("✅ İŞLEM TAMAMLANDI!")
        print(f"📄 Çıktı dosyası: {OUTPUT_FILE}")
        print("=" * 60)
        
    except Exception as e:
        print(f"\n❌ HATA: {str(e)}")
        import traceback
        traceback.print_exc()