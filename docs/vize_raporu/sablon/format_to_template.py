from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Dosya yolları (DOĞRU - boşluksuz)
BASE_DIR = r"D:\Projects\SelcukAiAssistant\docs\vize_raporu"
SOURCE_FILE = os.path.join(BASE_DIR, "VIZE_RAPORU.docx")  # Boşluk yok! 
TEMPLATE_FILE = os.path.join(BASE_DIR, "proje_sablonu.docx")  # Sablon ana dizinde
OUTPUT_FILE = os.path.join(BASE_DIR, "VIZE_RAPORU_FINAL.docx")

def check_files():
    """Dosyaların varlığını kontrol et"""
    print("🔍 Dosya kontrolü:")
    print(f"  Kaynak:  {SOURCE_FILE}")
    print(f"  Exists: {os.path.exists(SOURCE_FILE)}")
    print(f"  Şablon: {TEMPLATE_FILE}")
    print(f"  Exists: {os.path.exists(TEMPLATE_FILE)}")
    
    if not os.path.exists(SOURCE_FILE):
        print(f"\n❌ HATA:  Kaynak dosya bulunamadı!")
        return False
    
    if not os.path.exists(TEMPLATE_FILE):
        print(f"\n⚠️ Şablon bulunamadı, kaynak dosya üzerinden çalışılacak...")
        return "use_source"
    
    return True

def extract_sections(source_doc):
    """Kaynak dosyadan bölümleri çıkar"""
    doc = Document(source_doc)
    sections = {}
    current_section = None
    current_content = []
    
    print("📖 Bölümler çıkarılıyor...")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Bölüm başlıklarını tespit et
        if text in ['ÖZET', 'ABSTRACT', 'ÖNSÖZ', 'İÇİNDEKİLER', 
                    'SİMGELER VE KISALTMALAR', 'KAYNAKLAR']:
            if current_section and current_content: 
                sections[current_section] = current_content
            current_section = text
            current_content = []
            print(f"  ✓ {text}")
        
        # Numaralı bölümler (1. GİRİŞ, 2. KAYNAK ARAŞTIRMASI, vb.)
        elif text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            if current_section and current_content:
                sections[current_section] = current_content
            current_section = text
            current_content = []
            print(f"  ✓ {text}")
        
        # Alt başlıklar (1.1., 2.1., vb.)
        elif text and current_section: 
            current_content.append({
                'text': text,
                'style': para.style.name,
                'bold': para.runs[0].bold if para.runs else False
            })
    
    # Son bölümü ekle
    if current_section and current_content:
        sections[current_section] = current_content
    
    return sections

def fill_template(template_file, sections, output_file):
    """Şablonu sections ile doldur"""
    doc = Document(template_file)
    
    print("\n📝 Şablon dolduruluyor...")
    
    # Temel bilgileri değiştir
    replacements = {
        'PROJE BAŞLIĞINI BURAYA YAZINIZ': 'SELÇUK ÜNİVERSİTESİ İÇİN YAPAY ZEKA DESTEKLİ AKADEMİK ASİSTAN MOBİL UYGULAMASI',
        'Öğrencinin Adı SOYADI': 'Doğukan Balaman & Ali Yıldırım',
        'Unvanı Adı SOYADI': '[Danışman Unvanı ve Adı]',
        'Ay-Yıl': 'Aralık 2025',
        'Yıl, … Sayfa (Örnek:  2024, 105 Sayfa)': '2025, 85 Sayfa',
        '2024, 105 Sayfa':  '2025, 85 Sayfa',
    }
    
    # Paragrafları tara ve değiştir
    for para in doc.paragraphs:
        original_text = para.text
        for old, new in replacements.items():
            if old in para.text:
                # Inline replacement - formatı korur
                inline = para.runs
                for run in inline:
                    run.text = run.text.replace(old, new)
                print(f"  ✓ Değiştirildi: {old[: 30]}...")
        
        # Özet placeholder'ını doldur
        if 'Özet metnini yazmaya buradan başlayınız' in original_text:
            if 'ÖZET' in sections:
                para.clear()
                for item in sections['ÖZET']: 
                    if isinstance(item, dict):
                        run = para.add_run(item['text'] + '\n')
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                print("  ✓ ÖZET bölümü dolduruldu")
        
        # Abstract placeholder'ını doldur
        if 'Türkçe özet metninin İngilizce' in original_text:
            if 'ABSTRACT' in sections: 
                para.clear()
                for item in sections['ABSTRACT']: 
                    if isinstance(item, dict):
                        run = para.add_run(item['text'] + '\n')
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                print("  ✓ ABSTRACT bölümü dolduruldu")
    
    # Sayfa düzenini ayarla
    for section in doc.sections:
        section. left_margin = Cm(3.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
    
    doc.save(output_file)
    print(f"\n✅ Şablon kaydedildi: {output_file}")

def create_formatted_doc(sections, output_file):
    """Şablon yoksa sıfırdan oluştur"""
    doc = Document()
    
    print("\n📝 Yeni doküman oluşturuluyor...")
    
    # Sayfa düzeni
    section = doc.sections[0]
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    
    # Normal stil
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    
    # Bölümleri ekle
    for section_name, content in sections.items():
        # Başlık
        heading = doc. add_paragraph()
        run = heading.add_run(section_name)
        run.bold = True
        run.font. name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # İçerik
        for item in content:
            if isinstance(item, dict):
                para = doc.add_paragraph()
                run = para.add_run(item['text'])
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                if item.get('bold'):
                    run.bold = True
        
        doc.add_paragraph()  # Boşluk
    
    doc.save(output_file)
    print(f"✅ Doküman kaydedildi: {output_file}")

if __name__ == "__main__": 
    print("=" * 70)
    print("📄 SELÇUK ÜNİVERSİTESİ RAPOR FORMATLAMA ARACI")
    print("=" * 70)
    
    # Dosya kontrolü
    file_check = check_files()
    
    if file_check == False:
        print("\n❌ Gerekli dosyalar bulunamadı.  Çıkılıyor...")
        exit(1)
    
    try:
        # İçeriği çıkar
        sections = extract_sections(SOURCE_FILE)
        print(f"\n✅ Toplam {len(sections)} bölüm çıkarıldı\n")
        
        # Şablona göre işle
        if file_check == "use_source":
            create_formatted_doc(sections, OUTPUT_FILE)
        else:
            fill_template(TEMPLATE_FILE, sections, OUTPUT_FILE)
        
        print("\n" + "=" * 70)
        print("✅ İŞLEM BAŞARIYLA TAMAMLANDI!")
        print(f"📄 Çıktı dosyası: {OUTPUT_FILE}")
        print("=" * 70)
        
        # Dosyayı aç
        os.startfile(OUTPUT_FILE)
        
    except Exception as e:
        print(f"\n❌ HATA OLUŞTU: {str(e)}")
        import traceback
        traceback.print_exc()