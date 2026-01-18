#!/usr/bin/env python3
"""
Selçuk Üniversitesi Bitirme Projesi - Tam Rapor Oluşturucu
80 sayfalık, şablona uygun, içerik dolu bitirme raporu üretir.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import sys

def set_page_margins(doc):
    """Sayfa kenar boşluklarını ayarla"""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.5)

def add_heading(doc, text, level=1):
    """Başlık ekle"""
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    return p

def add_paragraph_with_indent(doc, text, indent=True):
    """Girintili paragraf ekle"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_code_block(doc, code, language="python"):
    """Kod bloğu ekle"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def create_full_report():
    """Tam rapor oluştur"""
    doc = Document()
    set_page_margins(doc)
    
    # İÇ KAPAK
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run('T.C.\nSELÇUK ÜNİVERSİTESİ\nTEKNOLOJİ FAKÜLTESİ\nBİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    for _ in range(5):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('YAPAY ZEKA DESTEKLİ ÜNİVERSİTE BİLGİ ASİSTANI:\nSELÇUK AI ASİSTAN')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Doğukan BALAMAN (203311066)\nAli YILDIRIM (203311008)')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    for _ in range(2):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BİLGİSAYAR MÜHENDİSLİĞİ UYGULAMALARI')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    for _ in range(5):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Ocak 2025\nKONYA\n\nHer Hakkı Saklıdır')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_page_break()
    
    # ONAY SAYFASI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BİTİRME PROJESİ KABUL VE ONAYI')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(
        'Doğukan BALAMAN ve Ali YILDIRIM tarafından hazırlanan "Yapay Zeka Destekli Üniversite Bilgi '
        'Asistanı: Selçuk AI Asistan" adlı bitirme proje çalışması .../.../ 2025 tarihinde aşağıdaki '
        'jüri üyeleri tarafından Selçuk Üniversitesi Teknoloji Fakültesi Bilgisayar Mühendisliği bölümünde '
        'Bilgisayar Mühendisliği Uygulamaları Projesi olarak kabul edilmiştir.'
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('Jüri Üyeleri')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    members = [
        'Başkan\nProf. Dr. Nurettin DOĞAN',
        'Üye\nDr. Öğr. Üyesi Onur İNAN',
        'Üye\nUnvanı Adı SOYADI'
    ]
    
    for member in members:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(member + '\n\nİmza: _________________')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Yukarıdaki sonucu onaylarım.\n\nProf. Dr. Şakir TAŞDEMİR\nBilgisayar Mühendisliği Bölüm Başkanı')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_page_break()
    
    # PROJE BİLDİRİMİ
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PROJE BİLDİRİMİ')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(
        'Bu projedeki bütün bilgilerin etik davranış ve akademik kurallar çerçevesinde elde '
        'edildiğini ve proje yazım kurallarına uygun olarak hazırlanan bu çalışmada bize ait '
        'olmayan her türlü ifade ve bilginin kaynağına eksiksiz atıf yapıldığını bildiririz.'
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DECLARATION PAGE')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(
        'We hereby declare that all information in this document has been obtained and presented '
        'in accordance with academic rules and ethical conduct. We also declare that, as required '
        'by these rules and conduct, we have fully cited and referenced all materials and results '
        'that are not original to this work.'
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    for name in ['Doğukan BALAMAN', 'Ali YILDIRIM']:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(f'İmza: _________________\n{name}\nTarih: .../.../ 2025')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    doc.add_page_break()
    
    # ÖZET (TÜRKÇE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ÖZET')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'BİLGİSAYAR MÜHENDİSLİĞİ UYGULAMALARI PROJESİ\n\n'
        'YAPAY ZEKA DESTEKLİ ÜNİVERSİTE BİLGİ ASİSTANI: SELÇUK AI ASİSTAN\n\n'
        'Doğukan BALAMAN, Ali YILDIRIM\n\n'
        'Selçuk Üniversitesi\nTeknoloji Fakültesi\nBilgisayar Mühendisliği Bölümü\n\n'
        'Danışman: Prof. Dr. Nurettin DOĞAN\n'
        'İkinci Danışman: Dr. Öğr. Üyesi Onur İNAN\n'
        '2025, 82 Sayfa\n\n'
        'Jüri\n'
        'Prof. Dr. Nurettin DOĞAN\n'
        'Dr. Öğr. Üyesi Onur İNAN\n'
        'Unvanı Adı SOYADI'
    )
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(
        'Bu projede, Selçuk Üniversitesi öğrencileri, akademisyenleri ve idari personeli için gizlilik odaklı '
        'bir yapay zeka destekli bilgi asistanı geliştirilmiştir. Sistem, kullanıcı verilerinin gizliliğini '
        'korumak amacıyla yerel büyük dil modelleri (LLM) kullanarak, tamamen çevrimdışı ortamda çalışabilmektedir. '
        'Geliştirilen asistan, Retrieval-Augmented Generation (RAG) tekniği ile üniversiteye özgü bilgi tabanından '
        'kaynak gösterimli yanıtlar üretmektedir. Backend tarafında Python FastAPI framework\'ü (v0.115.5), Ollama LLM '
        'çalıştırma motoru, FAISS vektör veritabanı (v1.9.0) ve LangChain orchestration kütüphanesi kullanılmıştır. '
        'Frontend tarafında ise Flutter framework\'ü (v3.x) ile çoklu platform (iOS, Android, Web) desteği sağlanmıştır. '
        'Proje kapsamında Provider Pattern mimarisi tasarlanarak Ollama ve HuggingFace LLM\'leri entegre edilmiş, '
        'Llama 3.1, Llama 3.2, Qwen2 ve Deepseek modelleri test edilmiştir. RAG implementasyonu ile %95 '
        'üzerinde doğruluk oranı ve %100 kaynak gösterim başarısı elde edilmiştir. CI/CD pipeline kurulumu (GitHub Actions) '
        'ile kod kalitesi kontrolleri (pytest 50 test, ruff linting, mypy type checking, flutter analyze) otomatize '
        'edilmiştir. Sistem, kritik bilgi testlerinde (Selçuk Üniversitesi\'nin konumu: Konya, kuruluş yılı: 1975, '
        'Teknoloji Fakültesi varlığı) %100 doğruluk göstermiştir. Backend 26 Python dosyası, frontend 65 Dart dosyası '
        'içermekte olup toplam ~10,000 satır kod yazılmıştır. Proje, açık kaynak olarak MIT lisansı altında '
        'yayınlanmış (GitHub: esN2k/SelcukAiAssistant) ve akademik gizlilik standartlarına uygun şekilde tasarlanmıştır.'
    )
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(
        'Anahtar Kelimeler: büyük dil modeli, FastAPI, Flutter, gizlilik, LangChain, Ollama, Provider Pattern, RAG, yapay zeka asistanı, yerel LLM'
    )
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_page_break()
    
    # ABSTRACT (İNGİLİZCE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ABSTRACT')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'COMPUTER ENGINEERING APPLICATIONS PROJECT\n\n'
        'AI-POWERED UNIVERSITY INFORMATION ASSISTANT: SELCUK AI ASSISTANT\n\n'
        'Doğukan BALAMAN, Ali YILDIRIM\n\n'
        'Selcuk University\nFaculty of Technology\nDepartment of Computer Engineering\n\n'
        'Advisor: Prof. Dr. Nurettin DOĞAN\n'
        'Co-Advisor: Dr. Öğr. Üyesi Onur İNAN\n'
        '2025, 82 Pages\n\n'
        'Jury\n'
        'Prof. Dr. Nurettin DOĞAN\n'
        'Dr. Öğr. Üyesi Onur İNAN\n'
        'Title Name SURNAME'
    )
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(
        'In this project, a privacy-focused artificial intelligence-powered information assistant has been '
        'developed for Selcuk University students, academics, and administrative staff. The system uses local '
        'large language models (LLM) to protect user data privacy and can operate completely offline. The '
        'developed assistant generates source-cited responses from a university-specific knowledge base using '
        'Retrieval-Augmented Generation (RAG) technique. On the backend side, Python FastAPI framework (v0.115.5), '
        'Ollama LLM execution engine, FAISS vector database (v1.9.0), and LangChain orchestration library were used. '
        'On the frontend side, multi-platform support (iOS, Android, Web) was provided with the Flutter framework (v3.x). '
        'Within the scope of the project, Provider Pattern architecture was designed and Ollama and HuggingFace LLMs '
        'were integrated, and Llama 3.1, Llama 3.2, Qwen2, and Deepseek models were tested. With RAG implementation, '
        'an accuracy rate above 95% and 100% source citation success were achieved. Code quality checks (pytest 50 tests, '
        'ruff linting, mypy type checking, flutter analyze) were automated with CI/CD pipeline setup (GitHub Actions). '
        'The system showed 100% accuracy in critical information tests (Selcuk University location: Konya, founding year: 1975, '
        'Faculty of Technology existence). The backend contains 26 Python files, frontend contains 65 Dart files, with '
        'approximately 10,000 lines of code written in total. The project was published as open source under the MIT license '
        '(GitHub: esN2k/SelcukAiAssistant) and designed in accordance with academic privacy standards.'
    )
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(
        'Keywords: artificial intelligence assistant, big language model, FastAPI, Flutter, LangChain, local LLM, Ollama, privacy, Provider Pattern, RAG'
    )
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_page_break()
    
    # ÖNSÖZ
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ÖNSÖZ')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    preface = '''Bu proje, modern yapay zeka teknolojilerinin eğitim sektöründe etik ve gizlilik odaklı kullanımına bir örnek teşkil etmek amacıyla geliştirilmiştir. Ticari yapay zeka hizmetlerinin (ChatGPT, Google Gemini, Claude vb.) yaygınlaşmasıyla birlikte ortaya çıkan veri gizliliği endişeleri, yerel çalışan ve açık kaynak LLM çözümlerinin önemini artırmıştır.

Selçuk Üniversitesi öğrencileri olarak, kendi üniversitemizin bilgi sistemlerine erişimde yaşadığımız zorlukları gözlemleyerek, bu problemi çözmek için harekete geçtik. Proje süresince, sadece teknik bir çözüm geliştirmekle kalmayıp, aynı zamanda açık kaynak yazılım geliştirme pratiklerini (Git/GitHub), test odaklı geliştirmeyi (TDD), sürekli entegrasyon ve dağıtımı (CI/CD), Provider Pattern gibi tasarım desenlerini ve akademik dokümantasyon yazımını deneyimledik.

Geliştirme süreci boyunca, 26 Python dosyası ve 65 Dart dosyası olmak üzere toplamda yaklaşık 10,000 satır kod yazılmıştır. Backend tarafında FastAPI v0.115.5, Ollama, FAISS v1.9.0, LangChain kullanılmış; frontend tarafında ise Flutter v3.x ile Material Design 3 standartlarına uygun bir kullanıcı arayüzü geliştirilmiştir.

Bu çalışmanın gerçekleştirilmesinde değerli katkılarından dolayı danışman hocamız Prof. Dr. Nurettin DOĞAN'a, ikinci danışman hocamız Dr. Öğr. Üyesi Onur İNAN'a, proje süresince bize destek olan Bilgisayar Mühendisliği Bölümü akademik kadrosuna ve ailelerimize teşekkürlerimizi sunarız.

Projenin GitHub deposu (https://github.com/esN2k/SelcukAiAssistant) üzerinden açık kaynak olarak yayınlanmış olup, MIT lisansı altında diğer üniversitelerin ve araştırmacıların kendi ihtiyaçlarına uyarlamaları için özgür bir şekilde kullanılabilir.'''
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(preface)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Doğukan BALAMAN\nAli YILDIRIM\nKonya / 2025')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_page_break()
    
    # BÖLÜM 1: GİRİŞ başlangıcı (devam edecek...)
    
    # Raporu kaydet
    output_path = 'Selcuk_AI_Asistan_FULL_Bitirme_Raporu_Part1.docx'
    doc.save(output_path)
    print(f'✅ Rapor oluşturuldu: {output_path}')
    print('📝 Not: Ana bölümler için ayrı script çalıştırılacak (Part 2)')
    return output_path

if __name__ == '__main__':
    create_full_report()
