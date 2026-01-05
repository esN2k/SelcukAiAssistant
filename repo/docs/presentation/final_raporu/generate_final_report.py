#!/usr/bin/env python3
"""
Selçuk Üniversitesi Bitirme Projesi Raporu Oluşturucu
Bu script, proje şablonuna uygun formatta Word dökümanı üretir.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def set_page_margins(doc):
    """Sayfa kenar boşluklarını ayarla: Sol 3.5cm, diğerleri 2.5cm"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.5)

def add_cover_page(doc):
    """İç kapak sayfası"""
    # T.C. SELÇUK ÜNİVERSİTESİ başlığı
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    run = p.add_run('T.C.\nSELÇUK ÜNİVERSİTESİ\nTEKNOLOJİ FAKÜLTESİ\nBİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    # Boşluk
    for _ in range(5):
        doc.add_paragraph()
    
    # Proje başlığı
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('YAPAY ZEKA DESTEKLİ ÜNİVERSİTE BİLGİ ASİSTANI:\nSELÇUK AI ASİSTAN')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    # Boşluk
    for _ in range(3):
        doc.add_paragraph()
    
    # Öğrenci isimleri
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Doğukan BALAMAN (203311066)\nAli YILDIRIM (203311008)')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    # Boşluk
    for _ in range(2):
        doc.add_paragraph()
    
    # Ders adı
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BİLGİSAYAR MÜHENDİSLİĞİ UYGULAMALARI')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    # Boşluk
    for _ in range(5):
        doc.add_paragraph()
    
    # Tarih ve yer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Ocak 2025\nKONYA')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    # Boşluk
    doc.add_paragraph()
    
    # Telif hakkı
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Her Hakkı Saklıdır')
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.italic = True
    
    doc.add_page_break()

def add_approval_page(doc):
    """Onay sayfası"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BİTİRME PROJESİ KABUL VE ONAYI')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(
        'Doğukan BALAMAN ve Ali YILDIRIM tarafından hazırlanan "Yapay Zeka Destekli Üniversite Bilgi '
        'Asistanı: Selçuk AI Asistan" adlı bitirme proje çalışması .../.../ 2025 tarihinde aşağıdaki '
        'jüri üyeleri tarafından Selçuk Üniversitesi Teknoloji Fakültesi Bilgisayar Mühendisliği bölümünde '
        'Bilgisayar Mühendisliği Uygulamaları Projesi olarak kabul edilmiştir.'
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Jüri üyeleri
    p = doc.add_paragraph()
    run = p.add_run('Jüri Üyeleri')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    members = [
        'Başkan\nProf. Dr. Nurettin DOĞAN',
        'Üye\nDr. Öğr. Üyesi Onur İNAN',
        'Üye\nUnvanı Adı SOYADI'
    ]
    
    for member in members:
        p = doc.add_paragraph()
        run = p.add_run(member + '\n\nİmza: _________________')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        doc.add_paragraph()
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Yukarıdaki sonucu onaylarım.')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Prof. Dr. Şakir TAŞDEMİR\nBilgisayar Mühendisliği Bölüm Başkanı')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_page_break()

def add_declaration_page(doc):
    """Proje bildirimi sayfası"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PROJE BİLDİRİMİ')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(
        'Bu projedeki bütün bilgilerin etik davranış ve akademik kurallar çerçevesinde elde '
        'edildiğini ve proje yazım kurallarına uygun olarak hazırlanan bu çalışmada bize ait '
        'olmayan her türlü ifade ve bilginin kaynağına eksiksiz atıf yapıldığını bildiririz.'
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DECLARATION PAGE')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(
        'We hereby declare that all information in this document has been obtained and presented '
        'in accordance with academic rules and ethical conduct. We also declare that, as required '
        'by these rules and conduct, we have fully cited and referenced all materials and results '
        'that are not original to this work.'
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # İmzalar
    for name in ['Doğukan BALAMAN', 'Ali YILDIRIM']:
        p = doc.add_paragraph()
        run = p.add_run(f'İmza: _________________\n{name}\nTarih: .../.../ 2025')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        doc.add_paragraph()
    
    doc.add_page_break()

def add_abstract_tr(doc):
    """Türkçe özet"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ÖZET')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    # Başlık bilgileri
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'BİLGİSAYAR MÜHENDİSLİĞİ UYGULAMALARI PROJESİ\n\n'
        'YAPAY ZEKA DESTEKLİ ÜNİVERSİTE BİLGİ ASİSTANI: SELÇUK AI ASİSTAN\n\n'
        'Doğukan BALAMAN, Ali YILDIRIM\n\n'
        'Selçuk Üniversitesi\nTeknoloji Fakültesi\nBilgisayar Mühendisliği Bölümü\n\n'
        'Danışman: Prof. Dr. Nurettin DOĞAN\n'
        'İkinci Danışman: Dr. Öğr. Üyesi Onur İNAN\n'
        '2025, 78 Sayfa\n\n'
        'Jüri\n'
        'Prof. Dr. Nurettin DOĞAN\n'
        'Dr. Öğr. Üyesi Onur İNAN\n'
        'Unvanı Adı SOYADI'
    )
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Özet metni
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(
        'Bu projede, Selçuk Üniversitesi öğrencileri, akademisyenleri ve idari personeli için gizlilik odaklı '
        'bir yapay zeka destekli bilgi asistanı geliştirilmiştir. Sistem, kullanıcı verilerinin gizliliğini '
        'korumak amacıyla yerel büyük dil modelleri (LLM) kullanarak, tamamen çevrimdışı ortamda çalışabilmektedir. '
        'Geliştirilen asistan, Retrieval-Augmented Generation (RAG) tekniği ile üniversiteye özgü bilgi tabanından '
        'kaynak gösterimli yanıtlar üretmektedir. Backend tarafında Python FastAPI framework\'ü, Ollama LLM '
        'çalıştırma motoru, FAISS vektör veritabanı ve LangChain orchestration kütüphanesi kullanılmıştır. '
        'Frontend tarafında ise Flutter framework\'ü ile çoklu platform (iOS, Android, Web) desteği sağlanmıştır. '
        'Proje kapsamında çoklu sağlayıcı (multi-provider) mimarisi tasarlanarak Ollama ve HuggingFace LLM\'leri '
        'entegre edilmiş, Llama 3.1, Qwen2 ve Deepseek modelleri test edilmiştir. RAG implementasyonu ile %95 '
        'üzerinde doğruluk oranı ve %100 kaynak gösterim başarısı elde edilmiştir. CI/CD pipeline kurulumu ile '
        'kod kalitesi kontrolleri (pytest, ruff, mypy, flutter analyze) otomatize edilmiştir. Sistem, kritik '
        'bilgi testlerinde (Selçuk Üniversitesi\'nin konumu, kuruluş yılı, fakülte bilgileri) %100 başarı '
        'göstermiştir. Proje, açık kaynak olarak MIT lisansı altında yayınlanmış ve akademik gizlilik '
        'standartlarına uygun şekilde tasarlanmıştır.'
    )
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Anahtar kelimeler
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(
        'Anahtar Kelimeler: büyük dil modeli, Flutter, gizlilik, Ollama, RAG, yapay zeka asistanı, yerel LLM'
    )
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_page_break()

def add_abstract_en(doc):
    """İngilizce özet"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ABSTRACT')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    # Başlık bilgileri
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'COMPUTER ENGINEERING APPLICATIONS PROJECT\n\n'
        'AI-POWERED UNIVERSITY INFORMATION ASSISTANT: SELCUK AI ASSISTANT\n\n'
        'Doğukan BALAMAN, Ali YILDIRIM\n\n'
        'Selcuk University\nFaculty of Technology\nDepartment of Computer Engineering\n\n'
        'Advisor: Prof. Dr. Nurettin DOĞAN\n'
        'Co-Advisor: Dr. Öğr. Üyesi Onur İNAN\n'
        '2025, 78 Pages\n\n'
        'Jury\n'
        'Prof. Dr. Nurettin DOĞAN\n'
        'Dr. Öğr. Üyesi Onur İNAN\n'
        'Title Name SURNAME'
    )
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Özet metni
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(
        'In this project, a privacy-focused artificial intelligence-powered information assistant has been '
        'developed for Selcuk University students, academics, and administrative staff. The system uses local '
        'large language models (LLM) to protect user data privacy and can operate completely offline. The '
        'developed assistant generates source-cited responses from a university-specific knowledge base using '
        'Retrieval-Augmented Generation (RAG) technique. On the backend side, Python FastAPI framework, Ollama '
        'LLM execution engine, FAISS vector database, and LangChain orchestration library were used. On the '
        'frontend side, multi-platform support (iOS, Android, Web) was provided with the Flutter framework. '
        'Within the scope of the project, a multi-provider architecture was designed and Ollama and HuggingFace '
        'LLMs were integrated, and Llama 3.1, Qwen2, and Deepseek models were tested. With RAG implementation, '
        'an accuracy rate above 95% and 100% source citation success were achieved. Code quality checks (pytest, '
        'ruff, mypy, flutter analyze) were automated with CI/CD pipeline setup. The system showed 100% success '
        'in critical information tests (location of Selcuk University, founding year, faculty information). The '
        'project was published as open source under the MIT license and designed in accordance with academic '
        'privacy standards.'
    )
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Anahtar kelimeler
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(
        'Keywords: artificial intelligence assistant, big language model, Flutter, local LLM, Ollama, privacy, RAG'
    )
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_page_break()

def add_preface(doc):
    """Önsöz"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ÖNSÖZ')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    preface_text = '''Bu proje, modern yapay zeka teknolojilerinin eğitim sektöründe etik ve gizlilik odaklı kullanımına bir örnek teşkil etmek amacıyla geliştirilmiştir. Ticari yapay zeka hizmetlerinin yaygınlaşmasıyla birlikte ortaya çıkan veri gizliliği endişeleri, yerel çalışan ve açık kaynak LLM çözümlerinin önemini artırmıştır.

Selçuk Üniversitesi öğrencileri olarak, kendi üniversitemizin bilgi sistemlerine erişimde yaşadığımız zorlukları gözlemleyerek, bu problemi çözmek için harekete geçtik. Proje süresince, sadece teknik bir çözüm geliştirmekle kalmayıp, aynı zamanda açık kaynak yazılım geliştirme pratiklerini, test odaklı geliştirmeyi (TDD), sürekli entegrasyon ve dağıtımı (CI/CD) deneyimledik.

Bu çalışmanın gerçekleştirilmesinde değerli katkılarından dolayı danışman hocamız Prof. Dr. Nurettin DOĞAN'a, ikinci danışman hocamız Dr. Öğr. Üyesi Onur İNAN'a, proje süresince bize destek olan Bilgisayar Mühendisliği Bölümü akademik kadrosuna ve ailelerimize teşekkürlerimizi sunarız.

Projenin GitHub deposu (https://github.com/esN2k/SelcukAiAssistant) üzerinden açık kaynak olarak yayınlanmış olup, diğer üniversitelerin ve araştırmacıların kendi ihtiyaçlarına uyarlamaları için özgür bir şekilde kullanılabilir.'''
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(preface_text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Doğukan BALAMAN\nAli YILDIRIM\nKonya / 2025')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_page_break()

def add_table_of_contents(doc):
    """İçindekiler"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('İÇİNDEKİLER')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    # İçindekiler listesi
    toc_items = [
        ('PROJE BİLDİRİMİ', 'iii'),
        ('ÖZET', 'iv'),
        ('ABSTRACT', 'v'),
        ('ÖNSÖZ', 'vi'),
        ('İÇİNDEKİLER', 'vii'),
        ('SİMGELER VE KISALTMALAR', 'ix'),
        ('', ''),
        ('1. GİRİŞ', '1'),
        ('2. KAYNAK ARAŞTIRMASI', '5'),
        ('3. MATERYAL VE YÖNTEM', '19'),
        ('4. SİSTEM TASARIMI VE UYGULAMA', '29'),
        ('5. ARAŞTIRMA BULGULARI VE TARTIŞMA', '46'),
        ('6. SONUÇLAR VE ÖNERİLER', '58'),
        ('', ''),
        ('KAYNAKLAR', '63'),
        ('EKLER', '67'),
        ('ÖZGEÇMİŞ', '77'),
    ]
    
    for item, page in toc_items:
        if item:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = p.add_run(f'{item}{"." * (80 - len(item) - len(page))}{page}')
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            if item.startswith(('1.', '2.', '3.', '4.', '5.', '6.', 'KAYNAKLAR', 'EKLER', 'ÖZGEÇMİŞ')):
                run.bold = True
        else:
            doc.add_paragraph()
    
    doc.add_page_break()

def add_abbreviations(doc):
    """Simgeler ve kısaltmalar"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SİMGELER VE KISALTMALAR')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('Simgeler')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run('Bu projede özel simge kullanılmamıştır.')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.italic = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('Kısaltmalar')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    abbreviations = [
        ('AI', 'Artificial Intelligence (Yapay Zeka)'),
        ('API', 'Application Programming Interface (Uygulama Programlama Arayüzü)'),
        ('CI/CD', 'Continuous Integration / Continuous Deployment (Sürekli Entegrasyon / Sürekli Dağıtım)'),
        ('CORS', 'Cross-Origin Resource Sharing (Çapraz Kaynak Paylaşımı)'),
        ('FAISS', 'Facebook AI Similarity Search (Facebook AI Benzerlik Araması)'),
        ('GPU', 'Graphics Processing Unit (Grafik İşleme Birimi)'),
        ('HF', 'HuggingFace'),
        ('HTTP', 'HyperText Transfer Protocol (Hipermetin Aktarım Protokolü)'),
        ('JSON', 'JavaScript Object Notation (JavaScript Nesne Gösterimi)'),
        ('KVKK', 'Kişisel Verilerin Korunması Kanunu'),
        ('LLM', 'Large Language Model (Büyük Dil Modeli)'),
        ('LoRA', 'Low-Rank Adaptation (Düşük Dereceli Adaptasyon)'),
        ('NLP', 'Natural Language Processing (Doğal Dil İşleme)'),
        ('OBS', 'Öğrenci Bilgi Sistemi'),
        ('RAG', 'Retrieval-Augmented Generation (Geri-Getirme Artırılmış Üretim)'),
        ('REST', 'Representational State Transfer (Temsili Durum Aktarımı)'),
        ('SSE', 'Server-Sent Events (Sunucu Taraflı Olaylar)'),
        ('TDD', 'Test-Driven Development (Test Odaklı Geliştirme)'),
        ('UI', 'User Interface (Kullanıcı Arayüzü)'),
        ('UTF-8', 'Unicode Transformation Format - 8 bit (Unicode Dönüşüm Formatı - 8 bit)'),
        ('UX', 'User Experience (Kullanıcı Deneyimi)'),
    ]
    
    for abbr, full in abbreviations:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(f'{abbr:<10}: {full}')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    doc.add_page_break()

def main():
    """Ana fonksiyon"""
    doc = Document()
    
    # Sayfa kenar boşluklarını ayarla
    set_page_margins(doc)
    
    # Ön kısım (Romen sayfa numaraları)
    add_cover_page(doc)
    add_approval_page(doc)
    add_declaration_page(doc)
    add_abstract_tr(doc)
    add_abstract_en(doc)
    add_preface(doc)
    add_table_of_contents(doc)
    add_abbreviations(doc)
    
    # Ana bölümler için ayrı bir dosya olacak (devam edecek)
    # Bu scriptin amacı şablonu göstermek
    
    # Dosyayı kaydet
    output_path = 'Selcuk_AI_Asistan_Bitirme_Raporu_Part1.docx'
    doc.save(output_path)
    print(f'Rapor oluşturuldu: {output_path}')
    print('Not: Ana bölümlerin geri kalanı ayrı olarak eklenecektir.')

if __name__ == '__main__':
    main()
