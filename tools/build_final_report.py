from __future__ import annotations

from pathlib import Path
from typing import Iterable
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "docs" / "presentation" / "final_raporu" / "Bilgisayar_Muhendisligi_Proje_Sablonu.docx"
OUTPUT = ROOT / "docs" / "presentation" / "final_raporu" / "Selcuk_AI_Asistan_Raporu.docx"
FIG_DIR = ROOT / "docs" / "presentation" / "final_raporu" / "figures"

TITLE_TR = "YAPAY ZEKA DESTEKLİ ÜNİVERSİTE BİLGİ ASİSTANI: SELÇUK AI ASİSTAN"
TITLE_EN = "AI-ASSISTED UNIVERSITY INFORMATION ASSISTANT: SELCUK AI ASSISTANT"
STUDENTS = ["Doğukan BALAMAN (203311066)", "Ali YILDIRIM (203311008)"]
ADVISORS = ["Prof. Dr. Nurettin DOĞAN", "Dr. Öğr. Üyesi Onur İNAN"]
DEPARTMENT = "BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ"
FACULTY = "TEKNOLOJİ FAKÜLTESİ"
UNIVERSITY = "SELÇUK ÜNİVERSİTESİ"
DATE_TEXT = "Ocak 2025"
CITY_TEXT = "KONYA"

BODY_STYLE = "Tez Metni_1.5 Satır"
BODY_STYLE_10 = "Tez Metni_1.0 Satır"
HEADING_1 = "Başlık 1. derece"
HEADING_2 = "Başlık 2. derece"
HEADING_3 = "Başlık 3. derece"
TABLE_CAPTION = "Çizelge Açıklaması"
FIG_CAPTION = "Şekil Açıklaması"


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    p.getparent().remove(p)


def find_paragraph(doc: Document, needle: str) -> Paragraph:
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise ValueError(f"Paragraph not found: {needle}")


def delete_between(doc: Document, start_text: str, end_text: str) -> Paragraph:
    start = find_paragraph(doc, start_text)
    end = find_paragraph(doc, end_text)
    deleting = False
    for p in list(doc.paragraphs):
        if p == start:
            deleting = True
            continue
        if p == end:
            break
        if deleting:
            delete_paragraph(p)
    return start


def add_heading_with_paragraphs(
    cursor: Paragraph,
    heading_text: str,
    level: int,
    paragraphs: Iterable[str],
    body_style: str = BODY_STYLE,
) -> Paragraph:
    style = {1: HEADING_1, 2: HEADING_2, 3: HEADING_3}[level]
    cursor = insert_paragraph_after(cursor, heading_text, style=style)
    for para in paragraphs:
        cursor = insert_paragraph_after(cursor, para, style=body_style)
    return cursor


def insert_table_after(paragraph: Paragraph, rows: int, cols: int):
    table = paragraph._parent.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    paragraph._p.addnext(table._tbl)
    return table


def set_table_font(table, size_pt: int = 10):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(size_pt)


def set_run_font(paragraph: Paragraph, size_pt: int, name: str | None = None):
    for run in paragraph.runs:
        run.font.size = Pt(size_pt)
        if name:
            run.font.name = name


def add_figure(
    cursor: Paragraph,
    image_path: Path,
    caption: str,
    width_inches: float = 5.5,
) -> Paragraph:
    cursor = insert_paragraph_after(cursor, "", style=BODY_STYLE)
    run = cursor.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    cursor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cursor = insert_paragraph_after(cursor, caption, style=FIG_CAPTION)
    cursor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return cursor


def _load_font(size: int) -> ImageFont.FreeTypeFont:
    for name in ["arial.ttf", "times.ttf", "calibri.ttf"]:
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_architecture_diagram(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1200, 700), "white")
    draw = ImageDraw.Draw(img)
    font = _load_font(22)

    boxes = {
        "Flutter UI": (80, 120, 360, 240),
        "FastAPI Backend": (420, 120, 760, 240),
        "Ollama LLM": (840, 60, 1130, 170),
        "HF LLM (ops.)": (840, 200, 1130, 310),
        "RAG (FAISS + Embedding)": (420, 320, 760, 450),
        "FAISS Index": (840, 380, 1130, 510),
    }

    for label, (x1, y1, x2, y2) in boxes.items():
        draw.rectangle([x1, y1, x2, y2], outline="black", width=3)
        w, h = draw.textbbox((0, 0), label, font=font)[2:]
        draw.text(((x1 + x2 - w) / 2, (y1 + y2 - h) / 2), label, fill="black", font=font)

    def arrow(start, end):
        draw.line([start, end], fill="black", width=3)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 10, ey - 6), (ex - 10, ey + 6)], fill="black")

    arrow((360, 180), (420, 180))
    arrow((760, 160), (840, 120))
    arrow((760, 200), (840, 250))
    arrow((590, 240), (590, 320))
    arrow((760, 385), (840, 445))

    title = "Selçuk AI Asistanı Genel Mimarisi"
    tw, th = draw.textbbox((0, 0), title, font=font)[2:]
    draw.text(((1200 - tw) / 2, 20), title, fill="black", font=font)

    img.save(path)


def draw_rag_pipeline(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1200, 600), "white")
    draw = ImageDraw.Draw(img)
    font = _load_font(20)

    steps = [
        ("Belge Toplama", 60),
        ("Temizleme", 260),
        ("Parçalama", 460),
        ("Gömme Üretimi", 660),
        ("FAISS İndeksi", 860),
    ]

    for label, x in steps:
        draw.rectangle([x, 220, x + 180, 320], outline="black", width=3)
        w, h = draw.textbbox((0, 0), label, font=font)[2:]
        draw.text((x + (180 - w) / 2, 245), label, fill="black", font=font)

    for i in range(len(steps) - 1):
        x1 = steps[i][1] + 180
        x2 = steps[i + 1][1]
        y = 270
        draw.line([(x1, y), (x2, y)], fill="black", width=3)
        draw.polygon([(x2, y), (x2 - 10, y - 6), (x2 - 10, y + 6)], fill="black")

    title = "RAG Süreci"
    tw, th = draw.textbbox((0, 0), title, font=font)[2:]
    draw.text(((1200 - tw) / 2, 40), title, fill="black", font=font)
    img.save(path)


def clean_doc_text(text: str) -> list[str]:
    lines = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        line = line.replace("\t", " ")
        line = re.sub(r"\s{2,}", " ", line)
        lines.append(line)
    return lines


abstract_tr = [
    "Bu çalışma, Selçuk Üniversitesi için geliştirilen yerel ve gizlilik odaklı yapay zeka destekli bilgi asistanının tasarımını ve uygulamasını sunmaktadır.",
    "Sistem, Flutter tabanlı arayüz, FastAPI tabanlı arka uç, yerel Ollama modeli ve FAISS tabanlı RAG katmanından oluşmaktadır.",
    "RAG katmanı, kaynaklı bağlam üretimiyle doğrulanabilir yanıtlar sağlamakta; SSE akışı ile gerçek zamanlı çıktı sunulmaktadır.",
    "Projede kritik bilgi doğruluk testleri ve CI kalite kontrolleri uygulanmış, sonuçlar akademik doğruluk ve gizlilik gereksinimlerini desteklemiştir.",
]
abstract_keywords_tr = "Anahtar Kelimeler: Yapay Zeka, RAG, FAISS, Yerel LLM, FastAPI, Flutter, Selçuk Üniversitesi"

abstract_en = [
    "This study presents a privacy‑focused local AI assistant developed for Selçuk University.",
    "The system combines a Flutter interface, a FastAPI backend, Ollama local models, and a FAISS‑based RAG layer.",
    "RAG provides source‑grounded answers, while SSE streaming delivers real‑time responses.",
    "Critical fact tests and CI quality checks support reliability and privacy requirements.",
]
abstract_keywords_en = "Keywords: Artificial Intelligence, RAG, FAISS, Local LLM, FastAPI, Flutter, Selcuk University"

preface_paragraphs = [
    "Bu proje, Selçuk Üniversitesi Teknoloji Fakültesi Bilgisayar Mühendisliği Bölümü kapsamında yürütülmüştür.",
    "Danışmanlarımız Prof. Dr. Nurettin DOĞAN ve Dr. Öğr. Üyesi Onur İNAN’a katkıları için teşekkür ederiz.",
]

symbols_paragraphs = [
    "LLM: Large Language Model (Büyük Dil Modeli)",
    "RAG: Retrieval-Augmented Generation (Geri Getirim Destekli Üretim)",
    "SSE: Server-Sent Events",
    "API: Application Programming Interface",
    "FAISS: Facebook AI Similarity Search",
    "HF: HuggingFace",
]

intro_paragraphs = [
    "Yapay zeka tabanlı bilgi asistanları, üniversite gibi büyük kurumsal yapılarda bilgiye erişimi hızlandırmaktadır.",
    "Transformer tabanlı LLM’ler, doğal dilde soru-cevap deneyimini yaygınlaştırmış; kurum içi veri gizliliği ihtiyacı yerel LLM çözümlerini ön plana çıkarmıştır.",
    "Bu proje, Selçuk Üniversitesi özelinde yerel LLM ve RAG birleşimini kullanarak kaynaklı ve güvenilir yanıt üretimini hedeflemiştir.",
]

importance_paragraphs = [
    "Yerel çalışma, kullanıcı verisinin kurum dışına çıkmasını önlemekte ve çevrimdışı kullanım senaryolarını desteklemektedir.",
    "RAG yaklaşımı, kaynaklı bağlam ile halüsinasyon riskini azaltmaktadır.",
]

scope_paragraphs = [
    "Kapsam; FastAPI arka uç servisleri, Ollama/HF sağlayıcıları, RAG indeksleme ve Flutter arayüz bileşenlerini kapsamaktadır.",
    "Kimlik doğrulama ve kurumsal SSO entegrasyonları bu çalışmanın kapsamı dışında bırakılmıştır.",
]

method_paragraphs = [
    "Geliştirme süreci iteratif şekilde yürütülmüş; dokümantasyon güncellemeleri ve kalite kontrolleri ile desteklenmiştir.",
    "Veri toplama için Selçuk Üniversitesi web sayfaları hedeflenmiş; scraping ve manuel doğrulama ile RAG veri seti hazırlanmıştır.",
]

system_paragraphs = [
    "Mimari, Flutter arayüzü ile FastAPI servisleri arasında HTTP/SSE iletişiminden oluşmaktadır.",
    "Model sağlayıcıları providers katmanında soyutlanmış; Ollama ve HuggingFace opsiyonel şekilde çalıştırılmaktadır.",
]

findings_paragraphs = [
    "pytest, ruff ve mypy tabanlı kalite kontrolleri CI iş akışında çalıştırılmaktadır.",
    "test_critical_facts.py ile Konya, 1975 ve Teknoloji Fakültesi gibi kritik bilgiler doğrulanmaktadır.",
]

conclusion_paragraphs = [
    "Sistem, yerel LLM ve RAG yaklaşımıyla gizlilik ve doğruluk gereksinimlerini karşılamaktadır.",
    "Gelecek çalışmalarda çok dilli destek ve sesli asistan özellikleri genişletilebilir.",
]

references = [
    "Vaswani, A., Shazeer, N., Parmar, N., et al., 2017, Attention Is All You Need, NeurIPS.",
    "Brown, T., Mann, B., Ryder, N., et al., 2020, Language Models are Few-Shot Learners, NeurIPS.",
    "Touvron, H., Martin, L., Stone, K., et al., 2023, LLaMA: Open and Efficient Foundation Language Models, arXiv.",
    "Lewis, P., Perez, E., Piktus, A., et al., 2020, Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks, NeurIPS.",
    "Johnson, J., Douze, M., ve Jégou, H., 2017, Billion-scale similarity search with GPUs, IEEE Transactions on Big Data.",
    "Reimers, N., ve Gurevych, I., 2019, Sentence-BERT, EMNLP.",
    "Ollama, 2024, Ollama Documentation, https://ollama.com (Erişim: 2025-12-30).",
    "FastAPI, 2024, FastAPI Documentation, https://fastapi.tiangolo.com (Erişim: 2025-12-30).",
    "Flutter, 2024, Flutter Documentation, https://docs.flutter.dev (Erişim: 2025-12-30).",
]


def main() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    arch_path = FIG_DIR / "architecture.png"
    rag_path = FIG_DIR / "rag_pipeline.png"
    draw_architecture_diagram(arch_path)
    draw_rag_pipeline(rag_path)

    doc = Document(str(TEMPLATE))

    # Cover table
    cover = doc.tables[0]
    cover_text_1 = f"T.C. {UNIVERSITY} {FACULTY} {DEPARTMENT}"
    cover_text_2 = (
        f"{TITLE_TR}\n"
        f"{STUDENTS[0]}\n"
        f"{STUDENTS[1]}\n"
        "MÜHENDİSLİK TASARIMI / BİLGİSAYAR MÜHENDİSLİĞİ UYGULAMALARI"
    )
    cover_text_3 = f"{DATE_TEXT} {CITY_TEXT} Her Hakkı Saklıdır"
    for col in range(2):
        cover.cell(1, col).text = cover_text_1
        cover.cell(2, col).text = cover_text_2
        cover.cell(3, col).text = cover_text_3

    # Approval paragraph
    approval_para = find_paragraph(doc, "tarafından hazırlan")
    approval_para.text = (
        f"{STUDENTS[0].split(' (')[0]} ve {STUDENTS[1].split(' (')[0]} tarafından hazırlanan "
        f\"{TITLE_TR}\" adlı proje çalışması …/…/2025 tarihinde aşağıdaki jüri üyeleri tarafından "
        "Selçuk Üniversitesi Teknoloji Fakültesi Bilgisayar Mühendisliği bölümünde Bitirme Projesi "
        "olarak kabul edilmiştir."
    )

    # Approval jury table
    jury_table = doc.tables[1]
    jury_table.cell(1, 0).text = f"Danışman {ADVISORS[0]}"
    jury_table.cell(2, 0).text = f"Üye {ADVISORS[1]} (İkinci Danışman)"
    jury_table.cell(3, 0).text = "Üye Unvanı Adı SOYADI"

    # Remove support footnote if present
    for p in list(doc.paragraphs):
        if p.text.strip().startswith("*Bu proje çalışması"):
            delete_paragraph(p)
        if "ikinci danışmanıdır" in p.text:
            p.text = f"**{ADVISORS[1]} bu proje çalışmasının ikinci danışmanıdır."

    # Declaration signature table
    decl_table = doc.tables[3]
    decl_table.cell(1, 1).text = f"{STUDENTS[0]}\n{STUDENTS[1]}"
    decl_table.cell(2, 1).text = "Tarih: …./…./2025"

    # Abstract (TR)
    find_paragraph(doc, "PROJE BAŞLIĞINI BURAYA YAZINIZ").text = TITLE_TR
    for p in doc.paragraphs:
        if p.text.strip() == "Öğrencinin Adı SOYADI":
            p.text = f"{STUDENTS[0]}\n{STUDENTS[1]}"
            break

    for p in doc.paragraphs:
        if p.text.strip().startswith("Danışman:"):
            p.text = f"Danışmanlar: {ADVISORS[0]}, {ADVISORS[1]}"
            break

    for p in doc.paragraphs:
        if p.text.strip().startswith("Yıl,"):
            p.text = "2025, … Sayfa"
            break

    abstract_anchor = find_paragraph(doc, "Özet metnini yazmaya buradan başlayınız")
    abstract_anchor.text = abstract_tr[0]
    set_run_font(abstract_anchor, 10)
    cursor = abstract_anchor
    for para in abstract_tr[1:]:
        cursor = insert_paragraph_after(cursor, para, style=BODY_STYLE_10)
        set_run_font(cursor, 10)

    keywords_tr = find_paragraph(doc, "Anahtar Kelimeler:")
    keywords_tr.text = abstract_keywords_tr
    set_run_font(keywords_tr, 10)

    # Abstract (EN)
    find_paragraph(doc, "PROJE BAŞLIĞININ İNGİLİZCE’SİNİ BURAYA YAZINIZ").text = TITLE_EN
    for p in doc.paragraphs:
        if p.text.strip().startswith("Year,"):
            p.text = "2025, … Pages"
            break

    abstract_en_anchor = find_paragraph(doc, "Türkçe özet metninin İngilizce’sini yazmaya buradan başlayınız")
    abstract_en_anchor.text = abstract_en[0]
    set_run_font(abstract_en_anchor, 10)
    cursor = abstract_en_anchor
    for para in abstract_en[1:]:
        cursor = insert_paragraph_after(cursor, para, style=BODY_STYLE_10)
        set_run_font(cursor, 10)

    keywords_en = find_paragraph(doc, "Keywords:")
    keywords_en.text = abstract_keywords_en
    set_run_font(keywords_en, 10)

    # Preface
    preface_anchor = find_paragraph(doc, "Önsöz metnini yazım kılavuzuna uygun olarak yazmaya buradan başlayınız")
    preface_anchor.text = preface_paragraphs[0]
    cursor = preface_anchor
    for para in preface_paragraphs[1:]:
        cursor = insert_paragraph_after(cursor, para, style=BODY_STYLE)

    preface_table = doc.tables[4]
    preface_table.cell(0, 1).text = f"{STUDENTS[0]}\n{STUDENTS[1]}"
    preface_table.cell(1, 1).text = "Konya / 2025"

    # Symbols & abbreviations
    symbols_anchor = find_paragraph(doc, "Simgeleri yazmaya buradan başlayınız")
    symbols_anchor.text = symbols_paragraphs[0]
    cursor = symbols_anchor
    for para in symbols_paragraphs[1:]:
        cursor = insert_paragraph_after(cursor, para, style=BODY_STYLE)

    for p in list(doc.paragraphs):
        if p.text.strip().startswith("Kısaltmaları yazmaya buradan başlayınız"):
            delete_paragraph(p)

    # Section 1: Giriş
    cursor = delete_between(doc, "1. GİRİŞ", "2. KAYNAK ARAŞTIRMASI")
    cursor = add_heading_with_paragraphs(cursor, "1.1. Projenin Arka Planı", 2, intro_paragraphs)
    cursor = add_heading_with_paragraphs(cursor, "1.2. Projenin Önemi", 2, importance_paragraphs)
    cursor = add_heading_with_paragraphs(cursor, "1.3. Projenin Kapsamı", 2, scope_paragraphs)

    # Section 2: Literature
    cursor = delete_between(doc, "2. KAYNAK ARAŞTIRMASI", "3. MATERYAL VE YÖNTEM")
    lit_doc = (ROOT / "docs" / "technical" / "ARCHITECTURE_OVERVIEW.md").read_text(encoding="utf-8", errors="ignore")
    lit_lines = clean_doc_text(lit_doc)
    cursor = add_heading_with_paragraphs(cursor, "2.1. Literatür Özeti", 2, lit_lines[:12])

    # Section 3: Material & Method
    cursor = delete_between(doc, "3. MATERYAL VE YÖNTEM", "4. ARAŞTIRMA SONUÇLARI VE TARTIŞMA")
    cursor = add_heading_with_paragraphs(cursor, "3.1. Geliştirme Yaklaşımı", 2, method_paragraphs)

    data_doc = (ROOT / "docs" / "reports" / "VERI_KAYNAKLARI.md").read_text(encoding="utf-8", errors="ignore")
    data_lines = clean_doc_text(data_doc)
    cursor = add_heading_with_paragraphs(cursor, "3.2. Veri Toplama", 2, data_lines[:10])

    # RAG params table
    table_caption = insert_paragraph_after(cursor, "Çizelge 3.1. RAG yapılandırma parametreleri", style=TABLE_CAPTION)
    cursor = table_caption
    table = insert_table_after(cursor, rows=5, cols=3)
    table.cell(0, 0).text = "Parametre"
    table.cell(0, 1).text = "Varsayılan"
    table.cell(0, 2).text = "Açıklama"
    table.cell(1, 0).text = "RAG_CHUNK_SIZE"
    table.cell(1, 1).text = "500"
    table.cell(1, 2).text = "Parça boyutu (karakter)"
    table.cell(2, 0).text = "RAG_CHUNK_OVERLAP"
    table.cell(2, 1).text = "50"
    table.cell(2, 2).text = "Parça bindirmesi"
    table.cell(3, 0).text = "RAG_TOP_K"
    table.cell(3, 1).text = "4"
    table.cell(3, 2).text = "Top-K parça sayısı"
    table.cell(4, 0).text = "RAG_EMBEDDING_MODEL"
    table.cell(4, 1).text = "paraphrase-multilingual-MiniLM-L12-v2"
    table.cell(4, 2).text = "Gömme modeli"
    set_table_font(table, 10)

    # Section 4: rename heading to System Design
    find_paragraph(doc, "4. ARAŞTIRMA SONUÇLARI VE TARTIŞMA").text = "4. SİSTEM TASARIMI VE UYGULAMA"
    cursor = delete_between(doc, "4. SİSTEM TASARIMI VE UYGULAMA", "5. SONUÇLAR VE ÖNERİLER")
    cursor = add_heading_with_paragraphs(cursor, "4.1. Genel Mimari", 2, system_paragraphs)
    cursor = add_figure(cursor, arch_path, "Şekil 4.1. Selçuk AI Asistanı genel mimarisi")

    api_doc = (ROOT / "docs" / "technical" / "API_CONTRACT.md").read_text(encoding="utf-8", errors="ignore")
    api_lines = clean_doc_text(api_doc)
    cursor = add_heading_with_paragraphs(cursor, "4.2. API Tasarımı", 2, api_lines[:12])

    cursor = add_figure(cursor, rag_path, "Şekil 4.2. RAG süreci")

    # Section 5: Findings
    find_paragraph(doc, "5. SONUÇLAR VE ÖNERİLER").text = "5. ARAŞTIRMA BULGULARI VE TARTIŞMA"
    cursor = delete_between(doc, "5. ARAŞTIRMA BULGULARI VE TARTIŞMA", "KAYNAKLAR")
    cursor = add_heading_with_paragraphs(cursor, "5.1. Test Bulguları", 2, findings_paragraphs)

    tests_doc = (ROOT / "docs" / "reports" / "TEST_RAPORU.md").read_text(encoding="utf-8", errors="ignore")
    tests_lines = clean_doc_text(tests_doc)
    cursor = add_heading_with_paragraphs(cursor, "5.2. CI ve Kalite Kontrolleri", 2, tests_lines[:12])

    # Insert new Section 6 before references
    ref_anchor = find_paragraph(doc, "KAYNAKLAR")
    cursor = insert_paragraph_after(ref_anchor, "6. SONUÇLAR VE ÖNERİLER", style=HEADING_1)
    cursor = add_heading_with_paragraphs(cursor, "6.1. Sonuçlar", 2, conclusion_paragraphs)

    # References
    delete_between(doc, "KAYNAKLAR", "EKLER")
    cursor = find_paragraph(doc, "KAYNAKLAR")
    for ref in references:
        cursor = insert_paragraph_after(cursor, ref, style=BODY_STYLE)

    # Appendices
    ek_anchor = find_paragraph(doc, "EKLER")
    cursor = ek_anchor

    cursor = add_heading_with_paragraphs(cursor, "EK-1: API Sözleşmesi (Özet)", 2, [])
    for line in api_lines:
        cursor = insert_paragraph_after(cursor, line, style=BODY_STYLE_10)
        set_run_font(cursor, 9, name="Courier New")

    cursor = add_heading_with_paragraphs(cursor, "EK-2: Örnek Kodlar", 2, [])
    code_files = [
        ROOT / "backend" / "main.py",
        ROOT / "backend" / "rag_service.py",
        ROOT / "backend" / "providers" / "ollama_provider.py",
        ROOT / "backend" / "providers" / "huggingface_provider.py",
        ROOT / "lib" / "controller" / "enhanced_chat_controller.dart",
    ]

    for code_path in code_files:
        cursor = add_heading_with_paragraphs(cursor, f"{code_path.as_posix()}", 3, [])
        code_text = code_path.read_text(encoding="utf-8", errors="ignore")
        for line in code_text.splitlines():
            cursor = insert_paragraph_after(cursor, line, style=BODY_STYLE_10)
            set_run_font(cursor, 9, name="Courier New")

    # CV section
    cv_anchor = find_paragraph(doc, "ÖZGEÇMİŞ")
    cursor = cv_anchor
    cursor = insert_paragraph_after(cursor, "ÖZGEÇMİŞ", style=HEADING_1)
    for student in STUDENTS:
        cursor = add_heading_with_paragraphs(cursor, student, 2, [])
        cursor = insert_paragraph_after(cursor, "Kişisel bilgiler, eğitim ve iletişim alanları öğrenci tarafından doldurulacaktır.", style=BODY_STYLE)

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
